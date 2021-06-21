import pandas as pd
from pandas.api.types import is_string_dtype
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math
import re


base = './EU DSA, DMA and consultations/Unzipped files/'
fp = 'contributions-Digital_Services_Act/contributions-Digital_Services_Act.xlsx'

DSA_contributions = pd.read_excel(base+fp).replace('&apos;', "'", regex=True)

pat = re.compile('^Please explain.{0,4}$')
column_names = DSA_contributions.columns.tolist()

for index, column_name in enumerate(DSA_contributions):
    if pat.search(column_name):
        
        column_names[index] = column_names[index-1] + ' Please explain.'

DSA_contributions.columns = column_names

letters = list('interoperability')
length = len(letters)

pats = [''.join(letters[0:n])+'..'+''.join(letters[n+2:length]) for n in range(length)]
pats = '|'.join(pats+['interoperable', 'interoperate'])

additional_cols = ['Reference', 'Organisation name', 'Country']

not_anonymous = DSA_contributions.dropna(subset=['Organisation name'])
in_english = not_anonymous[not_anonymous['Language'] == 'English']
filtered_DSA_contributions = in_english.sort_values('Organisation name')

question_frequencies = []
grouped = filtered_DSA_contributions.groupby('User type')

for user_type, contributions in grouped:
    all_interoperability = []
    
    for index, row in contributions.iterrows():
        test = row.str.contains(pats, na=False)
        as_list = []
        
        if test.any():
            for col in additional_cols:
                as_list.extend([(col, row[col])])
                
            filtered_row = row[test]
            as_list.extend(list(filtered_row.items()))
            
            all_interoperability.append(as_list)

    length = len(all_interoperability)

    as_doc = Document()
    
    style = as_doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    
    title = 'All responses mentioning "interoperability" (DSA consultation) - '
    title += user_type
    title += ' (n='+str(length)+')'
    
    as_doc.add_heading(title, 0)
    
    interop_questions = []

    for element in all_interoperability:
        new_questions = [q_a[0] for q_a in element[3:]]
        interop_questions.extend(new_questions)
        reference, organisation_name, country = element[0:3]
        
        heading_2 = reference[1]
        heading_2 += ' - '+organisation_name[1]
        heading_2 += ' ('+country[1]+')'
        
        as_doc.add_heading(heading_2, level=2)
        
        for question, content in element[3:]:
            as_doc.add_heading(question, level=3)
            split = content.splitlines()
            
            for text in split:
                words = re.split('(\W)', text)
                paragraph = as_doc.add_paragraph(words[0])
                
                for word in words[1:]:
                    if re.search(pats, word):
                        run = paragraph.add_run(word)
                        run.bold = True
                        
                    else:
                        paragraph.add_run(word)
                        
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
    frequencies = user_type, pd.Series(interop_questions, dtype='object').value_counts()
    question_frequencies.append(frequencies)
                
    fp = './EU DSA, DMA and consultations/Filtered contributions - NOT FOR ANNOTATING/'
    user_type = re.sub('/', '+', user_type)
    fp += 'DSA Interoperability responses - '+user_type+'.docx'
    
    as_doc.save(fp)